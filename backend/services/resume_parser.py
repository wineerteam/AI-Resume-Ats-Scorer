import io
import magic
from typing import Tuple, Optional, Tuple

import pdfplumber
from docx import Document
import PyPDF2

from backend.utils.file_utils import(
    FileParsingError, 
    TextExtractionError, 
    FileUploadError, 
    log_error, 
    log_warning, 
    log_info, 
    with_fallback
)

from backend.core.config import (
    MAX_FILE_SIZE_BYTES,
    MAX_FILE_SIZE_MB, 
    SUPPORTED_MIME_TYPES
)

class FileParsingError(Exception):
    pass

class FileValidationError(Exception):
    pass

def validate_file(file_data:bytes, filename:str)->Tuple[bool, str, Optional[str]]:
    file_size_bytes = len(file_data)
    if file_size_bytes > MAX_FILE_SIZE_BYTES:
        size_mb = file_size_bytes / (1024 * 1024)
        return False, (
            f'File size ({size_mb:.2f} MB) exceeds the maximum of {MAX_FILE_SIZE_MB} MB. '
            'Please upload a smaller file or compress your resume.'
        ), None
    
    if file_size_bytes==0:
        return False, 'uploade file is empty...please check the file you have uploaded and try again'
    
    try:
        mime_type = magic.from_buffer(file_data, mime=True)
    except Exception as e:
        from pathlib import Path
        ext = Path(filename).suffix.lower()
        if ext == '.pdf':
            mime_type = 'application/pdf'
        elif ext == '.docx':
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif ext in ['.doc', '.msword']:
            mime_type = 'application/msword'
        else:
            return False, f"error determining the file type: {e}", None
    
    if mime_type not in SUPPORTED_MIME_TYPES:
        supported=', '.join(SUPPORTED_MIME_TYPES.keys()).upper()
        return False, (
            f'Unsupported file type: {mime_type}. '
            f'Please upload one of: {supported}.'
        ), None
    
    

    return True, '', SUPPORTED_MIME_TYPES[mime_type]

def _extract_pdf_hyperlinks(file_data: bytes) -> str:
    urls = []
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_data))
        for page in reader.pages:
            if '/Annots' not in page:
                continue
            for annot_ref in page['/Annots']:
                try:
                    annot = annot_ref.get_object()
                    if annot.get('/Subtype') != '/Link':
                        continue
                    action = annot.get('/A', {})
                    uri = action.get('/URI', '')
                    if uri and isinstance(uri, (str, bytes)):
                        # PyPDF2 may return bytes for URI values
                        if isinstance(uri, bytes):
                            uri = uri.decode('utf-8', errors='ignore')
                        uri = uri.strip()
                        if uri.startswith('http'):
                            urls.append(uri)
                except Exception:
                    pass
    except Exception:
        pass
    return '\n'.join(urls)


def _extract_pdf_with_pdfplumber(file_data: bytes) -> str:
    text = ''
    with pdfplumber.open(io.BytesIO(file_data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + '\n'

    if not text.strip():
        raise TextExtractionError(
            'pdfplumber extracted no text',
            user_message='No text could be extracted from the PDF.'
        )
    
    hyperlinks = _extract_pdf_hyperlinks(file_data)
    if hyperlinks:
        text = text.strip() + '\n' + hyperlinks

    return text.strip()


def _extract_pdf_with_pypdf2(file_data: bytes) -> str:
    text = ''
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + '\n'

    if not text.strip():
        raise TextExtractionError(
            'PyPDF2 extracted no text',
            user_message='No text could be extracted from the PDF.'
        )

    hyperlinks = _extract_pdf_hyperlinks(file_data)
    if hyperlinks:
        text = text.strip() + '\n' + hyperlinks

    return text.strip()


def extract_text_from_pdf(file_data: bytes) -> str:
    try: 
        result, used_fallback=with_fallback(
        _extract_pdf_with_pdfplumber, 
        _extract_pdf_with_pypdf2, 
        file_data, 
        log_fallback=True
    )
    
        if used_fallback:
            log_info('PDF EXTRACTION succeded using the PyPDF2 fallback', context='resume_parser')
        return result
        
    except Exception as e:
        log_error(e, context='extract_text_from_pdf')
        raise FileParsingError(
            'Failed to extract text from PDF using both pdfplumber and PyPDF2. '
            'The PDF may be corrupted, password-protected, or contain only scanned images. '
            'Please ensure it contains selectable text.'
        ) from e
    

def extract_text_from_docx(file_data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_data))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        text = '\n'.join(text_parts)

        if not text.strip():
            raise FileParsingError(
                'No text could be extracted from the document. '
                'The document may be empty or corrupted.'
            )
        
        try:
            for rel in doc.part.rels.values():
                if 'hyperlink' in rel.reltype.lower():
                    url = rel._target
                    if isinstance(url, str) and url.startswith('http'):
                        text += '\n' + url
        except Exception:
            pass

        log_info(f'Extracted {len(text)} chars from DOCX', context='resume_parser')
        return text.strip()

    except FileParsingError:
        raise   # Re-raise unchanged — don't wrap in another FileParsingError

    except Exception as e:
        log_error(e, context='extract_text_from_docx')
        raise FileParsingError(
            'Failed to extract text from DOCX. '
            'The document may be corrupted or in an unsupported format. '
            'Please try re-saving or converting to PDF.'
        ) from e

def extract_text_from_doc(file_data: bytes) -> str:
    raise FileParsingError(
        'Legacy .doc format is not supported. '
        'Please convert your document to .docx or .pdf and try again. '
        'You can convert using Microsoft Word, Google Docs, or online tools.'
    )

def extract_text(file_data:bytes, file_type:str)->str:
    if file_type=='pdf':
        return extract_text_from_pdf(file_data)
    elif file_type=='docx':
        return extract_text_from_docx(file_data)
    elif file_type=='doc':
        return extract_text_from_doc(file_data)
    else:
        raise FileValidationError(
            f'invalid file type: {file_type}. supported types are: pdf, docx and doc'


        )
    
def parse_resume_file(file_data: bytes, filename:str)->Tuple[str, dict]:
    log_info(f'parsing file :{filename}', context='parse_Resume_file')

    #phase01:validate file
    try:
        is_valid, error_msg, file_type=validate_file(file_data, filename)
        if not is_valid:
            log_warning(f'valiudation failed for file {filename}', context='parse_resume_file')
            raise FileValidationError(error_msg)
    
    except FileValidationError as e:
        raise 

    except Exception as e:
        log_error(e, context='parse_resume_file_validation')
        raise FileValidationError(
            'Could not validate the uploaded file. Please ensure it is a valid PDF or DOCX.'
        ) from e
    
    #phase02: extraction of file

    try:
        text = extract_text(file_data, file_type)
        log_info(f'Extracted {len(text)} chars from {filename}', context='parse_resume_file')

    except FileParsingError:
        raise   # Re-raise unchanged

    except Exception as e:
        log_error(e, context='parse_resume_file_extraction')
        raise FileParsingError(
            'An unexpected error occurred while processing the file. '
            'Please try again or contact support if the problem persists.'
        ) from e

    metadata = {
        'filename':        filename,
        'file_type':       file_type,
        'file_size_bytes': len(file_data),
        'text_length':     len(text),
        'success':         True,
    }
    return text, metadata
